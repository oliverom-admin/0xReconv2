"""
Engagement Financial Impact Section Builder

Adds financial impact analysis to engagement executive summary reports.

Integrates financial calculations into the DOCX document with professional
formatting, risk cost visualization, and ROI analysis.

Usage:
    from engagement_financial_section import EngagementFinancialSectionBuilder

    # After building main content
    financial_builder = EngagementFinancialSectionBuilder()
    financial_builder.add_financial_section_to_document(
        doc,
        financial_summary=financial_data,
        style_manager=style_manager
    )
"""

import logging
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger('caip.reporting.engagement_financial_section')


class EngagementFinancialSectionBuilder:
    """
    Builds financial impact analysis section for engagement executive summary DOCX.

    Adds professional financial impact visualizations with executive-friendly
    interpretations and business case metrics.
    """

    def __init__(self):
        """Initialize financial section builder."""
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")

    def add_financial_section_to_document(self,
                                        doc: Document,
                                        financial_summary: Dict[str, Any],
                                        style_manager=None) -> None:
        """
        Add financial impact analysis section to DOCX document.

        Args:
            doc: python-docx Document object
            financial_summary: Output from EngagementFinancialCalculator.get_financial_summary()
            style_manager: Optional StyleManager for consistent styling
        """
        if not financial_summary:
            self.logger.info("No financial data to add")
            return

        try:
            self.logger.info("Adding financial impact section to document")

            # Add page break
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.add_page_break()

            # Add section heading
            heading = doc.add_heading('Financial Impact Analysis', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add introductory paragraph
            intro = doc.add_paragraph(
                'This section quantifies the financial risk of your current cryptographic '
                'asset exposure and the return on investment (ROI) for implementing recommended '
                'remediations. Financial estimates are based on industry benchmarks and '
                'conservative assumptions.\n\n'
                'The analysis measures risk in three dimensions:\n'
                '1. Annual Risk Cost - the estimated annual financial impact of current cryptographic vulnerabilities\n'
                '2. Remediation Investment - a phased 6-month programme to address identified issues\n'
                '3. Return on Investment - the annual savings and payback period for security improvements'
            )

            # Add annual risk cost section with page break
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.add_page_break()
            self._add_annual_risk_section(doc, financial_summary.get('annual_risk_cost', {}))

            # Add remediation costs section with page break
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.add_page_break()
            self._add_remediation_section(doc, financial_summary.get('remediation_costs', {}))

            # Add ROI analysis section with page break
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.add_page_break()
            self._add_roi_section(doc, financial_summary.get('roi_analysis', {}))

            # Add assumptions section with page break
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.add_page_break()
            self._add_assumptions_section(doc, financial_summary)

            self.logger.info("Successfully added financial section to document")

        except Exception as e:
            self.logger.error(f"Error adding financial section to document: {e}")
            raise

    def _add_annual_risk_section(self, doc: Document, annual_risk: Dict[str, Any]) -> None:
        """Add annual risk cost analysis section."""
        if not annual_risk:
            return

        try:
            doc.add_heading('Annual Risk Cost', level=2)

            # Main cost display
            total_cost = annual_risk.get('total_annual_cost', 0)
            risk_level = annual_risk.get('risk_level', 'UNKNOWN')

            # Add prominent cost figure
            cost_para = doc.add_paragraph()
            cost_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cost_run = cost_para.add_run(f"£{total_cost:,}")
            cost_run.font.size = Pt(28)
            cost_run.font.bold = True
            cost_run.font.color.rgb = self._get_risk_color(risk_level)

            risk_label = doc.add_paragraph(f"Annual Cost of Current Risk Exposure ({risk_level})")
            risk_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            risk_label.runs[0].font.size = Pt(11)

            # Breakdown table
            breakdown = annual_risk.get('breakdown', {})
            if breakdown:
                doc.add_paragraph()  # Spacing
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'

                header_cells = table.rows[0].cells
                header_cells[0].text = 'Risk Factor'
                header_cells[1].text = 'Count'

                risk_factors = [
                    ('Critical Findings', breakdown.get('critical_findings', 0)),
                    ('High Severity Findings', breakdown.get('high_findings', 0)),
                    ('Expired Certificates', breakdown.get('expired_certificates', 0)),
                    ('Weak Key Certificates', breakdown.get('weak_key_certificates', 0)),
                ]

                for factor, count in risk_factors:
                    row_cells = table.add_row().cells
                    row_cells[0].text = factor
                    row_cells[1].text = str(count)

            # Cost breakdown
            doc.add_paragraph()  # Spacing
            breach_cost = annual_risk.get('breach_risk_cost', 0)
            compliance_cost = annual_risk.get('compliance_risk_cost', 0)

            breakdown_table = doc.add_table(rows=1, cols=2)
            breakdown_table.style = 'Light Grid Accent 1'

            header_cells = breakdown_table.rows[0].cells
            header_cells[0].text = 'Cost Component'
            header_cells[1].text = 'Amount'

            row_cells = breakdown_table.add_row().cells
            row_cells[0].text = 'Breach Risk (Annual Probability)'
            row_cells[1].text = f"£{breach_cost:,}"

            row_cells = breakdown_table.add_row().cells
            row_cells[0].text = 'Compliance Risk (Regulatory)'
            row_cells[1].text = f"£{compliance_cost:,}"

            # Assumptions
            assumptions = annual_risk.get('assumptions', '')
            if assumptions:
                doc.add_paragraph()  # Spacing
                assumptions_para = doc.add_paragraph(f"Basis: {assumptions}")
                assumptions_para.runs[0].font.size = Pt(9)
                assumptions_para.runs[0].font.italic = True

        except Exception as e:
            self.logger.warning(f"Error adding annual risk section: {e}")

    def _add_remediation_section(self, doc: Document, remediation: Dict[str, Any]) -> None:
        """Add remediation costs and timeline section."""
        if not remediation:
            return

        try:
            doc.add_heading('Remediation Investment & Timeline', level=2)

            total_cost = remediation.get('total_cost', 0)
            phases = remediation.get('phases', [])

            # Total cost display
            total_para = doc.add_paragraph()
            total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            total_run = total_para.add_run(f"£{total_cost:,}")
            total_run.font.size = Pt(24)
            total_run.font.bold = True
            total_run.font.color.rgb = RGBColor(0, 100, 150)

            timeline_label = doc.add_paragraph("Total Investment Required")
            timeline_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timeline_label.runs[0].font.size = Pt(11)

            # Phase breakdown
            if phases:
                doc.add_paragraph()  # Spacing
                doc.add_paragraph("Phased Remediation Approach:").runs[0].font.bold = True

                for phase in phases:
                    phase_para = doc.add_paragraph(
                        f"{phase.get('name', 'Phase')}: £{phase.get('cost', 0):,}",
                        style='List Bullet'
                    )

                    duration = phase.get('duration_weeks', 0)
                    items = phase.get('items', '')
                    desc = phase.get('description', '')

                    details = f"{duration} weeks • {items}"
                    if desc:
                        details += f" • {desc}"

                    detail_para = doc.add_paragraph(details)
                    detail_para.paragraph_format.left_indent = Inches(0.5)
                    detail_para.runs[0].font.size = Pt(9)

        except Exception as e:
            self.logger.warning(f"Error adding remediation section: {e}")

    def _add_roi_section(self, doc: Document, roi_analysis: Dict[str, Any]) -> None:
        """Add ROI and payback analysis section."""
        if not roi_analysis:
            return

        try:
            doc.add_heading('Return on Investment (ROI)', level=2)

            investment = roi_analysis.get('remediation_investment', 0)
            annual_savings = roi_analysis.get('annual_risk_reduction', 0)
            roi_percent = roi_analysis.get('roi_percent', 0)
            payback_months = roi_analysis.get('payback_months', 0)
            three_year_savings = roi_analysis.get('roi_year3', 0)

            # ROI table
            roi_table = doc.add_table(rows=1, cols=2)
            roi_table.style = 'Light Grid Accent 1'

            header_cells = roi_table.rows[0].cells
            header_cells[0].text = 'Metric'
            header_cells[1].text = 'Value'

            metrics = [
                ('Annual Risk Reduction', f"£{annual_savings:,}"),
                ('Investment Required', f"£{investment:,}"),
                ('ROI Percentage', f"{roi_percent}%"),
                ('Payback Period', f"{payback_months} months"),
                ('3-Year Net Benefit', f"£{three_year_savings:,}"),
            ]

            for metric, value in metrics:
                row_cells = roi_table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value

            # Executive summary message
            message = roi_analysis.get('roi_message', '')
            if message:
                doc.add_paragraph()  # Spacing
                msg_para = doc.add_paragraph(message)
                msg_para.paragraph_format.space_before = Pt(12)
                msg_para.paragraph_format.space_after = Pt(12)
                msg_para.runs[0].font.bold = True

        except Exception as e:
            self.logger.warning(f"Error adding ROI section: {e}")

    def _add_assumptions_section(self, doc: Document, financial_summary: Dict[str, Any]) -> None:
        """Add disclaimer and assumptions section."""
        try:
            doc.add_heading('Assumptions & Disclaimer', level=2)

            assumptions_text = (
                'Financial estimates in this analysis are based on industry benchmarks and '
                'conservative assumptions. Actual costs may vary based on organizational factors, '
                'complexity of your environment, and implementation constraints.\n\n'
                'FINANCIAL MODEL EXPLANATION:\n'
                'Annual Risk Cost: Calculated by applying risk multipliers (based on finding severity) to the average '
                'data breach cost. A HIGH severity assessment (20% annual risk probability) results in an annual cost '
                'of 20% of the average breach cost. This represents the expected annual financial exposure.\n\n'
                'Remediation Investment: Estimated costs for addressing each identified finding or certificate issue. '
                'Costs are phased over 6 months (4/8/12 weeks per phase) to allow realistic implementation planning.\n\n'
                'ROI Analysis: Assumes 70% risk reduction from remediation (conservative). Annual savings = annual '
                'risk cost × 70%. Payback period shows months to recover investment from risk reduction.\n\n'
                'KEY BENCHMARKS (converted to GBP):\n'
                '• Average data breach cost: £3.6M (from IBM 2023 benchmark, $4.5M USD)\n'
                '• Cost per compromised record: £152 average\n'
                '• Average detection time: 207 days\n'
                '• Risk reduction from remediation: 70% (conservative estimate)\n'
                '• Compliance fine minimum: £21,000 per incident\n\n'
                'IMPORTANT: These estimates are for internal planning purposes only. Consult with your CFO, '
                'risk officer, and legal teams before making budget decisions. Actual breach costs can be significantly '
                'higher depending on industry, data sensitivity, customer impact, and regulatory jurisdiction.'
            )

            assumptions_para = doc.add_paragraph(assumptions_text)
            assumptions_para.runs[0].font.size = Pt(9)

            # Disclaimer box styling
            for run in assumptions_para.runs:
                run.font.color.rgb = RGBColor(100, 100, 100)

        except Exception as e:
            self.logger.warning(f"Error adding assumptions section: {e}")

    @staticmethod
    def _get_risk_color(risk_level: str) -> RGBColor:
        """Get RGB color for risk level."""
        colors = {
            'CRITICAL': RGBColor(198, 40, 40),    # Red
            'HIGH': RGBColor(245, 124, 0),        # Orange
            'MEDIUM': RGBColor(255, 167, 38),     # Amber
            'LOW': RGBColor(46, 125, 50),         # Green
        }
        return colors.get(risk_level, RGBColor(0, 0, 0))
