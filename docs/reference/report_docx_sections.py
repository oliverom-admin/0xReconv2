"""
Executive Summary Report Sections for DOCX - Phase 1

Provides data models and builders for executive summary content:
- Risk rating and posture
- Key metrics (findings by severity)
- Top action items
- Previous assessment tracking
- Organizational assessment summary

Data extraction from JSON assessment results.
Visualization-only layer (assessment is pre-calculated).

Usage:
    from report_docx_sections import ExecutiveSummaryBuilder, ExecutiveSummaryDataExtractor

    # Extract data from assessment JSON
    extractor = ExecutiveSummaryDataExtractor(scan_results)
    summary_data = extractor.extract()

    # Build section in document
    builder = ExecutiveSummaryBuilder()
    builder.build(document, style_manager, summary_data)
"""

from dataclasses import dataclass, field
from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .report_docx_styles import StyleManager, StyleColors


# ==================== Data Models ====================

@dataclass
class MetricBox:
    """Single metric for display (e.g., "237 Critical Findings")."""
    label: str
    value: str
    unit: str = ""
    severity: Optional[str] = None  # For color coding


@dataclass
class ActionItem:
    """Top action item with priority and details."""
    priority: int
    title: str
    description: str
    severity: str
    effort_estimate: Optional[str] = None
    owner: Optional[str] = None


@dataclass
class ExecutiveSummaryData:
    """
    Complete executive summary data extracted from assessment.

    All data comes from existing JSON assessment output.
    No new assessment calculations happen here.
    """
    # Risk assessment
    risk_rating: str  # CRITICAL, HIGH, MEDIUM, LOW
    risk_score: float  # 0-100 scale
    overall_health_index: float  # 0-100 scale

    # Key metrics (pre-calculated from assessment)
    total_findings: int
    critical_findings: int
    high_findings: int
    medium_findings: int
    low_findings: int
    informational_findings: int

    # Certificate inventory
    total_certificates: int
    expiring_certificates: int
    expired_certificates: int
    weak_key_algorithms: int
    non_pqc_ready: int

    # PQC readiness
    pqc_phase_1_count: int  # Ready for immediate migration
    pqc_phase_2_count: int  # Plan migration in 1-3 years
    pqc_phase_3_count: int  # Plan migration in 3-5 years
    pqc_phase_4_count: int  # Long-term or specialized use

    # Organizational assessment
    domain_1_score: float  # Organizational Readiness (%)
    domain_2_score: float  # Technology Constraints (%)
    domain_3_score: float  # Vendor Dependencies (%)

    # Previous assessment (if available)
    previous_assessment_date: Optional[str] = None
    previous_risk_rating: Optional[str] = None
    findings_trend: Optional[str] = None  # "improving", "stable", "declining"

    # Top action items (pre-identified in assessment)
    top_action_items: List[ActionItem] = field(default_factory=list)

    # Report metadata
    report_date: str = ""
    scan_date: Optional[str] = None
    engagement_name: str = ""


class ExecutiveSummaryDataExtractor:
    """
    Extract executive summary data from JSON assessment results.

    Reads pre-calculated assessment data and aggregates it into
    structures needed for professional executive presentation.

    Note: All calculations (risk scores, finding severity, etc.)
    are already done by the assessment engine. This just organizes
    the data for presentation.
    """

    def __init__(self, scan_results: Dict[str, Any]):
        """
        Initialize extractor with scan results JSON.

        Args:
            scan_results: Assessment JSON output from CAIP
        """
        self.scan_results = scan_results

    def extract(self) -> ExecutiveSummaryData:
        """
        Extract and aggregate executive summary data.

        Returns:
            ExecutiveSummaryData with all metrics populated
        """
        return ExecutiveSummaryData(
            # Risk assessment (from scoring)
            risk_rating=self._extract_risk_rating(),
            risk_score=self._extract_risk_score(),
            overall_health_index=self._extract_health_index(),

            # Findings breakdown (from policy assessment)
            total_findings=self._count_findings_total(),
            critical_findings=self._count_findings_by_severity('critical'),
            high_findings=self._count_findings_by_severity('high'),
            medium_findings=self._count_findings_by_severity('medium'),
            low_findings=self._count_findings_by_severity('low'),
            informational_findings=self._count_findings_by_severity('informational'),

            # Certificate inventory (from discovery)
            total_certificates=self._count_certificates(),
            expiring_certificates=self._count_expiring_certificates(),
            expired_certificates=self._count_expired_certificates(),
            weak_key_algorithms=self._count_weak_algorithms(),
            non_pqc_ready=self._count_non_pqc_ready(),

            # PQC readiness (from PQC analysis)
            pqc_phase_1_count=self._count_pqc_phase(1),
            pqc_phase_2_count=self._count_pqc_phase(2),
            pqc_phase_3_count=self._count_pqc_phase(3),
            pqc_phase_4_count=self._count_pqc_phase(4),

            # Organizational assessment (from assessment responses)
            domain_1_score=self._extract_domain_score('domain_1'),
            domain_2_score=self._extract_domain_score('domain_2'),
            domain_3_score=self._extract_domain_score('domain_3'),

            # Previous assessment tracking
            previous_assessment_date=self._extract_previous_date(),
            previous_risk_rating=self._extract_previous_rating(),
            findings_trend=self._calculate_trend(),

            # Top action items (pre-identified)
            top_action_items=self._extract_top_actions(),

            # Metadata
            report_date=datetime.now().strftime("%Y-%m-%d"),
            scan_date=self._extract_scan_date(),
            engagement_name=self.scan_results.get('engagement_name', 'Assessment'),
        )

    # ==================== Private Extraction Methods ====================

    def _extract_risk_rating(self) -> str:
        """Extract risk rating from assessment."""
        # Look for risk_rating in assessment results
        if 'assessment' in self.scan_results:
            rating = self.scan_results['assessment'].get('risk_rating', 'UNKNOWN')
            return rating.upper()
        return 'UNKNOWN'

    def _extract_risk_score(self) -> float:
        """Extract numerical risk score (0-100)."""
        if 'assessment' in self.scan_results:
            return self.scan_results['assessment'].get('risk_score', 0.0)
        return 0.0

    def _extract_health_index(self) -> float:
        """Extract overall health index."""
        if 'assessment' in self.scan_results:
            return self.scan_results['assessment'].get('health_index', 0.0)
        return 0.0

    def _count_findings_total(self) -> int:
        """Count total findings from policy assessment."""
        findings = self.scan_results.get('policy_assessment', {}).get('findings', [])
        return len(findings)

    def _count_findings_by_severity(self, severity: str) -> int:
        """Count findings by severity level."""
        findings = self.scan_results.get('policy_assessment', {}).get('findings', [])
        return sum(
            1 for f in findings
            if f.get('severity', '').lower() == severity.lower()
        )

    def _count_certificates(self) -> int:
        """Count total certificates in inventory."""
        inventory = self.scan_results.get('certificate_inventory', [])
        return len(inventory)

    def _count_expiring_certificates(self) -> int:
        """Count certificates expiring within 90 days."""
        inventory = self.scan_results.get('certificate_inventory', [])
        return sum(
            1 for cert in inventory
            if cert.get('days_to_expiry', 999) <= 90 and cert.get('days_to_expiry', 999) > 0
        )

    def _count_expired_certificates(self) -> int:
        """Count already-expired certificates."""
        inventory = self.scan_results.get('certificate_inventory', [])
        return sum(
            1 for cert in inventory
            if cert.get('days_to_expiry', 1) <= 0
        )

    def _count_weak_algorithms(self) -> int:
        """Count certificates using weak algorithms."""
        inventory = self.scan_results.get('certificate_inventory', [])
        weak_algs = ['md5', 'sha1', 'dsa', 'rsa-512', 'rsa-1024']
        return sum(
            1 for cert in inventory
            if any(weak in str(cert.get('signature_algorithm', '')).lower() for weak in weak_algs)
        )

    def _count_non_pqc_ready(self) -> int:
        """Count certificates not PQC-ready."""
        pqc_analysis = self.scan_results.get('pqc_analysis', {})
        assets = pqc_analysis.get('assets', [])
        return sum(
            1 for asset in assets
            if not asset.get('pqc_ready', False)
        )

    def _count_pqc_phase(self, phase: int) -> int:
        """Count assets in given PQC migration phase."""
        pqc_analysis = self.scan_results.get('pqc_analysis', {})
        assets = pqc_analysis.get('assets', [])
        return sum(
            1 for asset in assets
            if asset.get('migration_phase') == phase
        )

    def _extract_domain_score(self, domain: str) -> float:
        """Extract organizational assessment domain score."""
        assessment = self.scan_results.get('organizational_assessment', {})
        return assessment.get(f'{domain}_score', 0.0)

    def _extract_previous_date(self) -> Optional[str]:
        """Extract date of previous assessment."""
        history = self.scan_results.get('assessment_history', [])
        if history and len(history) > 1:
            return history[1].get('date')  # Second most recent
        return None

    def _extract_previous_rating(self) -> Optional[str]:
        """Extract risk rating from previous assessment."""
        history = self.scan_results.get('assessment_history', [])
        if history and len(history) > 1:
            return history[1].get('risk_rating')
        return None

    def _calculate_trend(self) -> Optional[str]:
        """Calculate trend (improving/stable/declining) from history."""
        history = self.scan_results.get('assessment_history', [])
        if not history or len(history) < 2:
            return None

        current_score = history[0].get('risk_score', 0)
        previous_score = history[1].get('risk_score', 0)

        if current_score < previous_score - 5:
            return 'improving'
        elif current_score > previous_score + 5:
            return 'declining'
        else:
            return 'stable'

    def _extract_top_actions(self) -> List[ActionItem]:
        """Extract top 3 action items from findings."""
        findings = self.scan_results.get('policy_assessment', {}).get('findings', [])

        # Sort by severity and priority score
        severity_order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}
        sorted_findings = sorted(
            findings,
            key=lambda f: (
                severity_order.get(f.get('severity', 'low').lower(), 4),
                -(f.get('priority_score', 0))
            )
        )

        # Take top 3
        items = []
        for i, finding in enumerate(sorted_findings[:3], 1):
            items.append(ActionItem(
                priority=i,
                title=finding.get('title', f'Finding {i}'),
                description=finding.get('description', ''),
                severity=finding.get('severity', 'medium'),
                effort_estimate=finding.get('effort_estimate'),
                owner=finding.get('recommended_owner'),
            ))

        return items

    def _extract_scan_date(self) -> Optional[str]:
        """Extract date of current scan."""
        return self.scan_results.get('scan_date', self.scan_results.get('timestamp'))


class ExecutiveSummaryBuilder:
    """
    Builds executive summary section in DOCX document.

    Creates professional visualization of pre-calculated assessment data:
    - Risk posture section
    - Key metrics tables
    - Top action items
    - Next steps checklist
    """

    @staticmethod
    def build(
        document: Document,
        style_manager: StyleManager,
        summary_data: ExecutiveSummaryData
    ) -> None:
        """
        Build complete executive summary section.

        Args:
            document: python-docx Document object
            style_manager: StyleManager instance
            summary_data: ExecutiveSummaryData with all metrics
        """
        # Section heading
        heading = document.add_paragraph()
        style_manager.apply_heading1(heading, "Executive Summary")

        # Risk posture box
        ExecutiveSummaryBuilder._build_risk_posture(
            document, style_manager, summary_data
        )

        # Key metrics
        ExecutiveSummaryBuilder._build_metrics_section(
            document, style_manager, summary_data
        )

        # Top action items
        if summary_data.top_action_items:
            ExecutiveSummaryBuilder._build_action_items(
                document, style_manager, summary_data.top_action_items
            )

        # Next steps
        ExecutiveSummaryBuilder._build_next_steps(
            document, style_manager
        )

        document.add_paragraph()

    @staticmethod
    def _build_risk_posture(
        document: Document,
        style_manager: StyleManager,
        summary_data: ExecutiveSummaryData
    ) -> None:
        """Build risk posture section with rating and key stats."""
        # Subheading
        heading = document.add_paragraph()
        style_manager.apply_heading2(heading, "Risk Posture")

        # Risk rating box
        risk_para = document.add_paragraph()
        risk_run = risk_para.add_run(f"Overall Risk Rating: ")
        risk_run.font.size = Pt(12)
        risk_run.font.bold = True

        rating_run = risk_para.add_run(summary_data.risk_rating)
        rating_run.font.size = Pt(14)
        rating_run.font.bold = True
        style_manager.apply_severity_text(rating_run, summary_data.risk_rating)

        # Brief stats
        stats_para = document.add_paragraph()
        stats_para.text = (
            f"{summary_data.total_findings} findings | "
            f"{summary_data.critical_findings} critical | "
            f"{summary_data.high_findings} high | "
            f"Health Index: {summary_data.overall_health_index:.0f}%"
        )
        stats_run = stats_para.runs[0] if stats_para.runs else stats_para.add_run()
        stats_run.font.size = Pt(11)
        style_manager.set_paragraph_spacing(stats_para, before=6, after=12)

    @staticmethod
    def _build_metrics_section(
        document: Document,
        style_manager: StyleManager,
        summary_data: ExecutiveSummaryData
    ) -> None:
        """Build key metrics table."""
        heading = document.add_paragraph()
        style_manager.apply_heading2(heading, "Key Metrics")

        # Create metrics table
        table = document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"

        # Data rows
        metrics = [
            ("Total Certificates", str(summary_data.total_certificates)),
            ("Expiring (90 days)", str(summary_data.expiring_certificates)),
            ("Weak Algorithms", str(summary_data.weak_key_algorithms)),
            ("PQC-Ready Assets", f"{summary_data.pqc_phase_1_count + summary_data.pqc_phase_2_count}"),
        ]

        for i, (label, value) in enumerate(metrics, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

        # Style table
        style_manager.style_table_alternating_rows(table)
        style_manager.add_table_borders(table)

        document.add_paragraph()

    @staticmethod
    def _build_action_items(
        document: Document,
        style_manager: StyleManager,
        items: List[ActionItem]
    ) -> None:
        """Build top action items section."""
        heading = document.add_paragraph()
        style_manager.apply_heading2(heading, "Top Action Items")

        for item in items:
            # Item heading
            item_heading = document.add_paragraph()
            item_heading.text = f"{item.priority}. {item.title}"
            item_run = item_heading.runs[0] if item_heading.runs else item_heading.add_run()
            item_run.font.bold = True
            item_run.font.size = Pt(11)
            style_manager.apply_severity_text(item_run, item.severity)

            # Item description
            if item.description:
                desc = document.add_paragraph(item.description)
                style_manager.apply_secondary_text(desc, item.description)

            # Details (effort, owner)
            if item.effort_estimate or item.owner:
                details = document.add_paragraph()
                if item.effort_estimate:
                    details.text += f"Effort: {item.effort_estimate}"
                if item.owner:
                    details.text += f" | Owner: {item.owner}"
                style_manager.apply_secondary_text(details, details.text)

            document.add_paragraph()

    @staticmethod
    def _build_next_steps(
        document: Document,
        style_manager: StyleManager
    ) -> None:
        """Build next steps checklist."""
        heading = document.add_paragraph()
        style_manager.apply_heading2(heading, "Next Steps")

        steps = [
            "Review and approve assessment findings with stakeholders",
            "Prioritize remediation actions based on risk and resources",
            "Assign owners for each critical/high finding",
            "Schedule quarterly reviews to track progress",
        ]

        for step in steps:
            para = document.add_paragraph(f"☐ {step}")
            style_manager.apply_body(para, f"☐ {step}")
