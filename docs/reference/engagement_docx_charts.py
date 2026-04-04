"""
Engagement DOCX Chart Section Builder

Adds visual charts to engagement executive summary reports.

Integrates chart PNG images into the DOCX document with professional
formatting, titles, and descriptions.

Usage:
    from engagement_docx_charts import EngagementChartSectionBuilder

    # After building main content
    chart_builder = EngagementChartSectionBuilder()
    chart_builder.add_charts_to_document(
        doc,
        charts={'certificate_inventory': '/path/to/chart.png', ...},
        style_manager=style_manager
    )
"""

import logging
import os
from typing import Dict, Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger('caip.reporting.engagement_docx_charts')


class EngagementChartSectionBuilder:
    """
    Builds charts section for engagement executive summary DOCX.

    Adds professional chart visualizations with titles and descriptions.
    """

    # Chart descriptions
    CHART_DESCRIPTIONS = {
        'certificate_inventory': (
            'Certificate Status Summary',
            'Shows the distribution of certificates by status. '
            'Valid certificates have >30 days until expiration. '
            'Expiring certificates expire within 30 days. '
            'Expired certificates are already past their expiration date.'
        ),
        'expiration_timeline': (
            'Certificate Expiration Timeline',
            'Visualizes the time to expiration for all certificates. '
            'Certificates expiring within 30 days require immediate attention. '
            'The 90+ day category represents certificates with sufficient time for renewal planning.'
        ),
        'algorithm_distribution': (
            'Certificate Algorithm Distribution',
            'Shows the most common signature algorithms across your certificate inventory. '
            'Modern algorithms (RSA-4096, ECDSA) provide stronger security than legacy algorithms. '
            'Legacy or non-standard algorithms should be prioritized for replacement.'
        ),
        'key_size_distribution': (
            'Key Size Distribution',
            'Breakdown of certificate key sizes by security strength. '
            'Weak keys (<2048 bits) no longer meet industry standards and pose security risk. '
            'Valid keys (2048-4096 bits) provide adequate protection. '
            'Strong keys (≥4096 bits) provide maximum protection for long-term data.'
        ),
        'finding_severity': (
            'Finding Severity Distribution',
            'Summary of policy assessment findings by severity level. '
            'Critical and High severity findings require immediate remediation. '
            'Medium findings should be addressed within 30 days. '
            'Low findings can be addressed through continuous improvement initiatives.'
        ),
    }

    def __init__(self):
        """Initialize chart section builder."""
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")

    def add_charts_to_document(self,
                              doc: Document,
                              charts: Dict[str, str],
                              style_manager=None,
                              charts_per_row: int = 1) -> None:
        """
        Add chart visualizations to DOCX document.

        Args:
            doc: python-docx Document object
            charts: Dictionary mapping chart names to file paths
            style_manager: Optional StyleManager for consistent styling
            charts_per_row: Number of charts per row (1 or 2)
        """
        if not charts:
            self.logger.info("No charts to add")
            return

        try:
            self.logger.info(f"Adding {len(charts)} charts to document")

            # Add section heading
            heading = doc.add_heading('Visual Analytics', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add introductory paragraph
            intro = doc.add_paragraph(
                'The following charts provide visual insights into your cryptographic asset '
                'inventory, certificate health, and policy compliance status. Each visualization '
                'highlights key risks and operational metrics across your organization.'
            )

            # Add each chart
            for chart_name, chart_path in sorted(charts.items()):
                if not os.path.exists(chart_path):
                    self.logger.warning(f"Chart file not found: {chart_path}")
                    continue

                # Get chart info
                title, description = self.CHART_DESCRIPTIONS.get(
                    chart_name,
                    (chart_name.replace('_', ' ').title(), 'Chart visualization.')
                )

                # Add page break before chart (unless it's the first)
                if doc.paragraphs and doc.paragraphs[-1].text:
                    doc.add_page_break()

                # Add chart title
                chart_heading = doc.add_heading(title, level=2)
                chart_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add chart image
                try:
                    # Add picture with optimal width
                    doc.add_picture(chart_path, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.logger.debug(f"Added chart: {chart_name}")
                except Exception as e:
                    self.logger.error(f"Error adding chart image {chart_name}: {e}")
                    doc.add_paragraph(f"[Chart image unavailable: {chart_name}]")

                # Add description
                desc_para = doc.add_paragraph(description)
                desc_para.paragraph_format.space_before = Pt(6)
                desc_para.paragraph_format.space_after = Pt(12)
                desc_para.runs[0].font.size = Pt(10)

            self.logger.info(f"Successfully added charts to document")

        except Exception as e:
            self.logger.error(f"Error adding charts to document: {e}")
            raise

    @staticmethod
    def add_simple_charts(doc: Document, charts: Dict[str, str]) -> None:
        """
        Simplified method to add charts with minimal formatting.

        Args:
            doc: python-docx Document object
            charts: Dictionary mapping chart names to file paths
        """
        if not charts:
            return

        for chart_name, chart_path in sorted(charts.items()):
            if os.path.exists(chart_path):
                try:
                    title = chart_name.replace('_', ' ').title()
                    doc.add_heading(title, level=2)
                    doc.add_picture(chart_path, width=Inches(5.5))
                    doc.add_paragraph()  # Spacing
                except Exception as e:
                    logger.warning(f"Could not add chart {chart_name}: {e}")
