"""
DOCX Styling System for Executive Reports - Phase 1

Provides centralized style management for professional Word document generation.
Implements corporate color palette and consistent typography for all report sections.

Color Scheme:
- Primary: Navy Blue (#0D1B2A) - Headers, main text
- Secondary: Steel Blue (#415A77) - Subheadings, borders
- Severity colors: Critical (red), High (orange), Medium (yellow), Low (teal)

Typography:
- Headings: Calibri, bold, navy
- Body: Calibri, 11pt, dark slate
- Emphasis: Bold navy for key metrics
- Line spacing: 1.15x (professional readability)
"""

from typing import Optional, Tuple
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


class StyleColors:
    """
    Centralized color palette for executive report styling.
    Ensures consistency across all document sections.

    All colors use RGB tuples for compatibility with python-docx RGBColor.
    """

    # Primary colors (navy-based corporate palette)
    PRIMARY = RGBColor(13, 27, 42)           # Deep navy #0D1B2A
    PRIMARY_LIGHT = RGBColor(27, 58, 95)    # Navy blue #1B3A5F
    SECONDARY = RGBColor(65, 90, 119)       # Steel blue #415A77
    ACCENT = RGBColor(119, 141, 169)        # Muted blue #778DA9

    # Severity colors
    CRITICAL = RGBColor(193, 18, 31)        # Deep red #C1121F
    CRITICAL_BG = RGBColor(255, 229, 229)   # Light red #FFE5E5

    HIGH = RGBColor(227, 100, 20)           # Burnt orange #E36414
    HIGH_BG = RGBColor(255, 243, 224)       # Light orange #FFF3E0

    MEDIUM = RGBColor(244, 162, 97)         # Sandy orange #F4A261
    MEDIUM_BG = RGBColor(255, 248, 225)     # Light yellow #FFF8E1

    LOW = RGBColor(42, 157, 143)            # Teal #2A9D8F
    LOW_BG = RGBColor(224, 242, 241)        # Light teal #E0F2F1

    INFO = RGBColor(108, 117, 125)          # Gray #6C757D
    INFO_BG = RGBColor(245, 245, 245)       # Light gray #F5F5F5

    # Text colors
    TEXT_PRIMARY = RGBColor(30, 41, 59)     # Dark slate #1E293B
    TEXT_SECONDARY = RGBColor(100, 116, 139) # Slate gray #64748B
    TEXT_LIGHT = RGBColor(255, 255, 255)    # White

    # Background colors
    BG_PRIMARY = RGBColor(248, 250, 252)    # Very light #F8FAFC
    BG_WHITE = RGBColor(255, 255, 255)      # Pure white

    # Borders
    BORDER = RGBColor(203, 213, 225)        # Border gray #CBD5E1

    # Other
    SUCCESS = RGBColor(5, 150, 105)         # Green #059669


class StyleManager:
    """
    Centralized style application for DOCX report generation.

    Handles:
    - Font sizing and styling (headings, body, emphasis)
    - Color application (text and background)
    - Table styling (alternating rows, borders)
    - Paragraph spacing and alignment

    Usage:
        style_mgr = StyleManager()
        style_mgr.apply_heading1(paragraph, "Section Title")
        style_mgr.apply_body(paragraph, "Body text...")
        style_mgr.style_table_alternating_rows(table)
    """

    def __init__(self):
        """Initialize style manager with default settings."""
        self.colors = StyleColors()

    # ==================== Heading Styles ====================

    def apply_heading1(self, paragraph, text: str) -> None:
        """
        Apply Heading 1 style: 24pt, bold, navy, 18pt spacing after.
        Used for major section headers (e.g., "Executive Summary", "Key Findings").

        Args:
            paragraph: python-docx paragraph object
            text: Heading text to apply style to
        """
        paragraph.text = text
        paragraph.style = 'Heading 1'

        # Configure run properties
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = self.colors.PRIMARY

        # Paragraph spacing
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15

    def apply_heading2(self, paragraph, text: str) -> None:
        """
        Apply Heading 2 style: 16pt, bold, secondary blue, 12pt spacing after.
        Used for subsection headers.

        Args:
            paragraph: python-docx paragraph object
            text: Heading text to apply style to
        """
        paragraph.text = text
        paragraph.style = 'Heading 2'

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = self.colors.SECONDARY

        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15

    def apply_heading3(self, paragraph, text: str) -> None:
        """
        Apply Heading 3 style: 13pt, bold, accent blue, 8pt spacing after.
        Used for sub-subsections.

        Args:
            paragraph: python-docx paragraph object
            text: Heading text to apply style to
        """
        paragraph.text = text
        paragraph.style = 'Heading 3'

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = self.colors.ACCENT

        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15

    # ==================== Body Text Styles ====================

    def apply_body(self, paragraph, text: str) -> None:
        """
        Apply standard body text style: 11pt, dark slate, 1.15 line spacing.
        Used for all regular content paragraphs.

        Args:
            paragraph: python-docx paragraph object
            text: Body text to apply style to
        """
        paragraph.text = text

        if paragraph.runs:
            run = paragraph.runs[0]
        else:
            run = paragraph.add_run()

        run.font.size = Pt(11)
        run.font.color.rgb = self.colors.TEXT_PRIMARY

        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(6)

    def apply_emphasis(self, run) -> None:
        """
        Apply emphasis style to a run: bold, navy color.
        Used for key metrics, important terms.

        Args:
            run: python-docx run object
        """
        run.font.bold = True
        run.font.color.rgb = self.colors.PRIMARY

    def apply_secondary_text(self, paragraph, text: str) -> None:
        """
        Apply secondary text style: 10pt, slate gray.
        Used for captions, notes, supplementary information.

        Args:
            paragraph: python-docx paragraph object
            text: Secondary text
        """
        paragraph.text = text

        if paragraph.runs:
            run = paragraph.runs[0]
        else:
            run = paragraph.add_run()

        run.font.size = Pt(10)
        run.font.color.rgb = self.colors.TEXT_SECONDARY

        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(4)

    # ==================== Severity and Status Styles ====================

    def get_severity_color(self, severity: str) -> RGBColor:
        """
        Get color for a severity level.

        Args:
            severity: One of 'critical', 'high', 'medium', 'low', 'info'

        Returns:
            RGBColor for the severity level
        """
        severity_lower = severity.lower()

        if severity_lower == 'critical':
            return self.colors.CRITICAL
        elif severity_lower == 'high':
            return self.colors.HIGH
        elif severity_lower == 'medium':
            return self.colors.MEDIUM
        elif severity_lower == 'low':
            return self.colors.LOW
        else:
            return self.colors.INFO

    def get_severity_bg_color(self, severity: str) -> RGBColor:
        """
        Get background color for a severity level (light variant).

        Args:
            severity: One of 'critical', 'high', 'medium', 'low', 'info'

        Returns:
            Light RGBColor for the severity level
        """
        severity_lower = severity.lower()

        if severity_lower == 'critical':
            return self.colors.CRITICAL_BG
        elif severity_lower == 'high':
            return self.colors.HIGH_BG
        elif severity_lower == 'medium':
            return self.colors.MEDIUM_BG
        elif severity_lower == 'low':
            return self.colors.LOW_BG
        else:
            return self.colors.INFO_BG

    def apply_severity_text(self, run, severity: str) -> None:
        """
        Apply severity color to a run.

        Args:
            run: python-docx run object
            severity: Severity level
        """
        run.font.color.rgb = self.get_severity_color(severity)
        run.font.bold = True

    # ==================== Table Styles ====================

    def style_table_alternating_rows(self, table) -> None:
        """
        Apply alternating row styling to a table.
        - Header row: navy background, white text, bold
        - Odd rows: light background
        - Even rows: white background
        - All rows: navy borders

        Args:
            table: python-docx Table object
        """
        tbl = table._element
        tblPr = tbl.tblPr

        # Create table style
        tblStyle = parse_xml(
            f'<w:tblStyle {nsdecls("w")} w:val="TableGrid"/>'
        )

        if tblPr.find(qn('w:tblStyle')) is not None:
            tblPr.remove(tblPr.find(qn('w:tblStyle')))
        tblPr.append(tblStyle)

        # Apply row styling
        for i, row in enumerate(table.rows):
            # Header row (first row)
            if i == 0:
                self._style_header_row(row)
            # Alternating body rows
            else:
                if i % 2 == 0:
                    self._style_light_row(row)
                else:
                    self._style_white_row(row)

    def _style_header_row(self, row) -> None:
        """Style a table header row."""
        for cell in row.cells:
            # Background color
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="0D1B2A"/>'
            )
            cell._element.get_or_add_tcPr().append(shading_elm)

            # Text styling
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = self.colors.TEXT_LIGHT
                    run.font.size = Pt(11)

    def _style_light_row(self, row) -> None:
        """Style a light alternating row."""
        for cell in row.cells:
            # Light background
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>'
            )
            cell._element.get_or_add_tcPr().append(shading_elm)

    def _style_white_row(self, row) -> None:
        """Style a white alternating row."""
        for cell in row.cells:
            # White background (explicit)
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="FFFFFF"/>'
            )
            cell._element.get_or_add_tcPr().append(shading_elm)

    def add_table_borders(self, table) -> None:
        """
        Add consistent borders to all table cells.

        Args:
            table: python-docx Table object
        """
        tbl = table._element
        tblPr = tbl.tblPr

        # Table borders
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CBD5E1')  # Border gray
            tblBorders.append(border)

        tblPr.append(tblBorders)

    # ==================== Utility Methods ====================

    def set_paragraph_spacing(self, paragraph, before: float = 6, after: float = 6) -> None:
        """
        Set consistent paragraph spacing (in points).

        Args:
            paragraph: python-docx paragraph object
            before: Points before paragraph
            after: Points after paragraph
        """
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15

    def center_align(self, paragraph) -> None:
        """Center-align a paragraph."""
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def right_align(self, paragraph) -> None:
        """Right-align a paragraph."""
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def justify_align(self, paragraph) -> None:
        """Justify-align a paragraph."""
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def parse_xml(xml_string: str):
    """Helper function to parse XML strings."""
    from docx.oxml import parse_xml as docx_parse_xml
    return docx_parse_xml(xml_string)
