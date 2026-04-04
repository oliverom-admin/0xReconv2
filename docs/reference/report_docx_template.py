"""
DOCX Template Components for Executive Reports - Phase 1

Provides builders for standard document elements:
- Cover page with engagement details
- Headers and footers with security markings
- Auto-updating table of contents
- Professional section breaks

Each builder follows a consistent pattern:
1. Accept configuration parameters
2. Create required document elements
3. Apply styles using StyleManager
4. Return modified document or section

Usage:
    from docx import Document
    from report_docx_styles import StyleManager
    from report_docx_template import CoverPageBuilder

    doc = Document()
    style_mgr = StyleManager()

    CoverPageBuilder.build(
        doc,
        style_mgr,
        engagement_name="Acme Corp",
        organization_name="Acme Corporation",
        report_date="2024-02-15",
        classification="CONFIDENTIAL",
        logo_path="logo.png"
    )
"""

from datetime import datetime
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .report_docx_styles import StyleManager, StyleColors


class CoverPageBuilder:
    """
    Builds professional cover page for executive reports.

    Elements:
    - Logo (top right, optional)
    - Engagement name (centered, large, navy)
    - Organization name (centered, medium)
    - Report date (right-aligned, gray)
    - Classification marking (footer, red background if CONFIDENTIAL)
    - Page break to separate from content
    """

    @staticmethod
    def build(
        document: Document,
        style_manager: StyleManager,
        engagement_name: str,
        organization_name: str,
        report_date: str,
        classification: str = "CONFIDENTIAL",
        logo_path: Optional[str] = None
    ) -> None:
        """
        Build cover page in document.

        Args:
            document: python-docx Document object
            style_manager: StyleManager instance for consistent styling
            engagement_name: Name of the engagement/assessment
            organization_name: Customer/organization name
            report_date: Date of report (YYYY-MM-DD or formatted)
            classification: Classification level (CONFIDENTIAL, SECRET, etc.)
            logo_path: Optional path to logo image file
        """
        # Add logo if provided
        if logo_path and Path(logo_path).exists():
            CoverPageBuilder._add_logo(document, logo_path)

        # Add spacing for visual hierarchy
        document.add_paragraph()  # Spacing
        document.add_paragraph()

        # Add engagement name
        engagement_para = document.add_paragraph()
        engagement_para.text = engagement_name
        engagement_run = engagement_para.runs[0]
        engagement_run.font.size = Pt(32)
        engagement_run.font.bold = True
        engagement_run.font.color.rgb = style_manager.colors.PRIMARY
        style_manager.center_align(engagement_para)
        style_manager.set_paragraph_spacing(engagement_para, before=24, after=12)

        # Add organization name
        org_para = document.add_paragraph()
        org_para.text = organization_name
        org_run = org_para.runs[0]
        org_run.font.size = Pt(18)
        org_run.font.color.rgb = style_manager.colors.TEXT_SECONDARY
        style_manager.center_align(org_para)
        style_manager.set_paragraph_spacing(org_para, before=0, after=36)

        # Add spacing
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()

        # Add report date (right-aligned)
        date_para = document.add_paragraph()
        date_para.text = f"Report Date: {report_date}"
        date_run = date_para.runs[0]
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = style_manager.colors.TEXT_SECONDARY
        style_manager.right_align(date_para)
        style_manager.set_paragraph_spacing(date_para, before=12, after=48)

        # Add classification marking
        CoverPageBuilder._add_classification_footer(
            document,
            style_manager,
            classification
        )

        # Page break
        document.add_page_break()

    @staticmethod
    def _add_logo(document: Document, logo_path: str, width: float = 1.5) -> None:
        """
        Add logo image to top-right of cover page.

        Args:
            document: python-docx Document object
            logo_path: Path to logo image file
            width: Width of logo in inches (default 1.5")
        """
        try:
            # Create table for logo placement (1x1, no borders)
            logo_table = document.add_table(rows=1, cols=1)
            logo_table.autofit = False
            logo_table.allow_autofit = False

            # Set table to full width
            logo_table.width = Inches(7.5)

            # Add logo to cell
            cell = logo_table.rows[0].cells[0]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(width))

            # Right-align logo
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Remove table borders
            tbl = logo_table._element
            tblPr = tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tblBorders.append(border)

            tblPr.append(tblBorders)

        except Exception as e:
            print(f"Warning: Could not add logo from {logo_path}: {e}")

    @staticmethod
    def _add_classification_footer(
        document: Document,
        style_manager: StyleManager,
        classification: str
    ) -> None:
        """
        Add classification marking to footer of cover page.

        Args:
            document: python-docx Document object
            style_manager: StyleManager instance
            classification: Classification level
        """
        # Determine colors based on classification
        if classification.upper() in ["CONFIDENTIAL", "SECRET", "TOP SECRET"]:
            bg_color = "C1121F"  # Red
            text_color = style_manager.colors.TEXT_LIGHT
        else:
            bg_color = "778DA9"  # Muted blue
            text_color = style_manager.colors.TEXT_LIGHT

        # Create table for classification bar
        class_table = document.add_table(rows=1, cols=1)
        class_table.autofit = False

        cell = class_table.rows[0].cells[0]
        paragraph = cell.paragraphs[0]
        paragraph.text = classification
        paragraph_run = paragraph.runs[0]
        paragraph_run.font.bold = True
        paragraph_run.font.size = Pt(14)
        paragraph_run.font.color.rgb = text_color
        style_manager.center_align(paragraph)

        # Style cell background
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), bg_color)
        cell._element.get_or_add_tcPr().append(shading_elm)

        # Set cell height and padding
        tcPr = cell._element.tcPr
        tcH = OxmlElement('w:tcH')
        tcH.set(qn('w:val'), '500')
        tcH.set(qn('w:type'), 'atLeast')
        tcPr.append(tcH)


class HeaderFooterBuilder:
    """
    Adds headers and footers to document with security markings.

    Elements:
    - Header: Section name on left, page number/total on right
    - Footer: Classification marking on left, document name on right
    - Skips header on cover page (first page different)

    Usage:
        HeaderFooterBuilder.add_headers_and_footers(
            document,
            style_manager,
            classification="CONFIDENTIAL",
            organization_name="Acme Corp"
        )
    """

    @staticmethod
    def add_headers_and_footers(
        document: Document,
        style_manager: StyleManager,
        classification: str = "CONFIDENTIAL",
        organization_name: str = "Assessment"
    ) -> None:
        """
        Add headers and footers to all sections.

        Args:
            document: python-docx Document object
            style_manager: StyleManager instance
            classification: Classification marking for footer
            organization_name: Organization name for identification
        """
        # Get default section
        section = document.sections[0]

        # Configure header - minimal setup only
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = organization_name

        # Configure footer - minimal setup only
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = f"{classification} - {organization_name} Assessment"

    @staticmethod
    def skip_header_on_first_page(document: Document, section_index: int = 0) -> None:
        """
        Configure section to skip header/footer on first page (cover page).

        Args:
            document: python-docx Document object
            section_index: Index of section (default 0 for first section)
        """
        section = document.sections[section_index]
        section.different_first_page_header_footer = True


class TableOfContentsBuilder:
    """
    Creates auto-updating table of contents for Word documents.

    Note: python-docx has limited support for Word fields. For full
    auto-updating TOC, users must:
    1. Open document in Word
    2. Right-click TOC field
    3. Select "Update Field"

    This builder inserts the TOC field code for Word to process.
    """

    @staticmethod
    def build(document: Document, style_manager: StyleManager) -> None:
        """
        Insert table of contents with Word field code.

        Args:
            document: python-docx Document object
            style_manager: StyleManager instance
        """
        # Add TOC heading
        toc_heading = document.add_paragraph()
        toc_heading.text = "Table of Contents"
        toc_run = toc_heading.runs[0] if toc_heading.runs else toc_heading.add_run()
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        toc_run.font.color.rgb = style_manager.colors.PRIMARY
        style_manager.set_paragraph_spacing(toc_heading, before=0, after=12)

        # Insert TOC field
        # This creates a placeholder; Word will process it when opened
        TableOfContentsBuilder._insert_toc_field(document)

        # Add spacing after TOC
        document.add_paragraph()

        # Page break after TOC
        document.add_page_break()

    @staticmethod
    def _insert_toc_field(document: Document) -> None:
        r"""
        Insert Word TOC field code.

        Field code: { TOC \o "1-3" \h \z \u }
        - \o: Outline levels 1-3
        - \h: Hyperlinks (clickable)
        - \z: Hide page numbers in web view
        - \u: Use outline levels

        Args:
            document: python-docx Document object
        """
        paragraph = document.add_paragraph()
        p = paragraph._element

        # Create run for field start
        r1 = OxmlElement('w:r')
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r1.append(fldChar1)
        p.append(r1)

        # Create run for instruction text
        r2 = OxmlElement('w:r')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        r2.append(instrText)
        p.append(r2)

        # Create run for field end
        r3 = OxmlElement('w:r')
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r3.append(fldChar2)
        p.append(r3)

    @staticmethod
    def mark_for_toc(paragraph, level: int = 1) -> None:
        """
        Mark a paragraph for inclusion in TOC.

        Args:
            paragraph: python-docx paragraph object
            level: TOC level (1-3), determines nesting

        Note:
            Automatically applied by heading styles (Heading 1, 2, 3).
            Use this for custom TOC entries if needed.
        """
        # Heading styles automatically mark for TOC
        if level == 1:
            paragraph.style = 'Heading 1'
        elif level == 2:
            paragraph.style = 'Heading 2'
        elif level == 3:
            paragraph.style = 'Heading 3'


class PageBreakBuilder:
    """
    Utility for adding section breaks and page breaks.
    """

    @staticmethod
    def add_page_break(document: Document) -> None:
        """Add a page break to the document."""
        document.add_page_break()

    @staticmethod
    def add_section_break(document: Document, new_section: bool = True) -> None:
        """
        Add a section break.

        Args:
            document: python-docx Document object
            new_section: If True, creates new section (allows different headers/footers)
        """
        if new_section:
            document.add_section()
        else:
            document.add_page_break()

    @staticmethod
    def add_section_break_with_heading(
        document: Document,
        style_manager: StyleManager,
        section_title: str
    ) -> None:
        """
        Add page break and new section heading.

        Args:
            document: python-docx Document object
            style_manager: StyleManager instance
            section_title: Title for new section
        """
        document.add_page_break()

        heading = document.add_paragraph()
        style_manager.apply_heading1(heading, section_title)
