"""
ExecutiveDocxService — DOCX executive report generation using python-docx.

Sections: cover → executive summary → findings → remediation → technical → financial → appendix
"""
from __future__ import annotations

from datetime import datetime
from typing import Any

import structlog
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = structlog.get_logger("recon.executive_docx")

# Station Hex colour scheme
COLOR_ACCENT = RGBColor(0x00, 0xFF, 0x41)
COLOR_SECONDARY = RGBColor(0xFF, 0xB8, 0x00)
COLOR_DARK = RGBColor(0x0D, 0x1B, 0x2A)
COLOR_CRITICAL = RGBColor(0xFF, 0x44, 0x44)
COLOR_HIGH = RGBColor(0xFF, 0x88, 0x00)
COLOR_MEDIUM = RGBColor(0xFF, 0xCC, 0x00)
COLOR_LOW = RGBColor(0x44, 0xAA, 0xFF)
COLOR_INFO = RGBColor(0x88, 0x88, 0x88)

SEVERITY_COLORS = {
    "critical": COLOR_CRITICAL,
    "high": COLOR_HIGH,
    "medium": COLOR_MEDIUM,
    "low": COLOR_LOW,
    "info": COLOR_INFO,
}


class ExecutiveDocxService:
    def __init__(self, report_data: dict[str, Any]) -> None:
        self._data = report_data

    def generate(self, output_path: str) -> str:
        """Generate DOCX and write to output_path. Returns path."""
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)

        self._add_cover_page(doc)
        self._add_executive_summary(doc)
        self._add_findings_breakdown(doc)
        self._add_remediation_roadmap(doc)
        self._add_technical_details(doc)
        self._add_financial_impact(doc)
        self._add_appendix(doc)

        doc.save(output_path)
        logger.info("docx_generated", path=output_path)
        return output_path

    def _add_cover_page(self, doc: Document) -> None:
        doc.add_paragraph("")
        doc.add_paragraph("")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("0xRecon")
        run.font.size = Pt(36)
        run.font.color.rgb = COLOR_ACCENT
        run.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Executive Report")
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_DARK

        doc.add_paragraph("")
        for label, value in [
            ("Project", self._data.get("project_name", "—")),
            ("Report", self._data.get("report_name", "—")),
            ("Scan", self._data.get("scan_name", "—")),
            ("Generated", self._data.get("generated_at", "—")[:19]),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(12)
            r2 = p.add_run(value)
            r2.font.size = Pt(12)

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document) -> None:
        doc.add_heading("Executive Summary", level=1)
        summary = self._data.get("summary") or {}
        grade = summary.get("grade", "N/A")
        score = summary.get("health_score", 0)

        p = doc.add_paragraph()
        p.add_run(f"Health Score: {score:.1f}/100  |  Grade: {grade}").bold = True

        doc.add_paragraph(
            f"This assessment identified {summary.get('total_findings', 0)} findings "
            f"across {summary.get('total_certificates', 0)} certificates and "
            f"{summary.get('total_keys', 0)} keys."
        )

        # Severity table
        sev = summary.get("findings_by_severity") or {}
        if sev:
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "Severity"
            hdr[1].text = "Count"
            for cell in hdr:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            for severity in ["critical", "high", "medium", "low", "info"]:
                count = sev.get(severity, 0)
                if count > 0:
                    row = table.add_row().cells
                    row[0].text = severity.upper()
                    row[1].text = str(count)
                    color = SEVERITY_COLORS.get(severity)
                    if color:
                        for p in row[0].paragraphs:
                            for r in p.runs:
                                r.font.color.rgb = color

    def _add_findings_breakdown(self, doc: Document) -> None:
        doc.add_heading("Findings Breakdown", level=1)
        findings = self._data.get("findings") or []

        if not findings:
            doc.add_paragraph("No findings were identified during this assessment.")
            return

        # Group by severity
        for severity in ["critical", "high", "medium", "low", "info"]:
            sev_findings = [
                f for f in findings
                if (f.get("severity") or "info").lower() == severity
            ]
            if not sev_findings:
                continue

            h = doc.add_heading(f"{severity.upper()} ({len(sev_findings)})", level=2)
            color = SEVERITY_COLORS.get(severity)
            if color:
                for r in h.runs:
                    r.font.color.rgb = color

            for f in sev_findings:
                title = f.get("title") or f.get("rule_name") or "Finding"
                p = doc.add_paragraph()
                p.add_run(title).bold = True
                if f.get("entity_cn"):
                    p.add_run(f"  ({f['entity_cn']})")

                if f.get("description"):
                    doc.add_paragraph(f["description"], style="List Bullet")
                if f.get("remediation"):
                    rem_p = doc.add_paragraph(style="List Bullet")
                    rem_p.add_run("Remediation: ").bold = True
                    rem_p.add_run(f["remediation"])

    def _add_remediation_roadmap(self, doc: Document) -> None:
        doc.add_heading("Remediation Roadmap", level=1)
        findings = self._data.get("findings") or []

        if not findings:
            doc.add_paragraph("No remediation actions required.")
            return

        # Group by urgency
        immediate = [f for f in findings
                     if (f.get("severity") or "").lower() in ("critical",)]
        short_term = [f for f in findings
                      if (f.get("severity") or "").lower() in ("high",)]
        planned = [f for f in findings
                   if (f.get("severity") or "").lower() in ("medium", "low", "info")]

        for group_name, group in [
            ("Immediate Action (Critical)", immediate),
            ("Short-term (High Priority)", short_term),
            ("Planned Remediation", planned),
        ]:
            if not group:
                continue
            doc.add_heading(group_name, level=2)
            for i, f in enumerate(group, 1):
                p = doc.add_paragraph(style="List Number")
                title = f.get("title") or f.get("rule_name") or "Action item"
                p.add_run(title).bold = True
                if f.get("remediation"):
                    doc.add_paragraph(f"    {f['remediation']}")

    def _add_technical_details(self, doc: Document) -> None:
        doc.add_heading("Technical Details", level=1)

        certs = self._data.get("certificates") or []
        keys = self._data.get("keys") or []

        doc.add_paragraph(
            f"Certificate Inventory: {len(certs)} certificates discovered"
        )
        doc.add_paragraph(f"Key Inventory: {len(keys)} keys discovered")

        if certs:
            doc.add_heading("Certificate Summary", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, text in enumerate(["Subject CN", "Algorithm", "Key Size", "Status"]):
                hdr[i].text = text
                for p in hdr[i].paragraphs:
                    for r in p.runs:
                        r.bold = True

            for cert in certs[:20]:  # Limit to 20 in executive report
                row = table.add_row().cells
                cn = cert.get("subject_cn") or cert.get("subject", {}).get("CN", "—")
                row[0].text = str(cn)[:40]
                row[1].text = cert.get("key_algorithm") or cert.get("public_key_algorithm") or "—"
                row[2].text = str(cert.get("key_size") or cert.get("public_key_size") or "—")
                expired = cert.get("is_expired", False)
                row[3].text = "Expired" if expired else "Valid"

    def _add_financial_impact(self, doc: Document) -> None:
        fin = self._data.get("financial_impact")
        if not fin:
            return

        doc.add_heading("Financial Impact Analysis", level=1)

        arc = fin.get("annual_risk_cost") or {}
        if arc:
            doc.add_heading("Annual Risk Cost", level=2)
            doc.add_paragraph(
                f"Risk Level: {arc.get('risk_level', '—')}"
            )
            cost = arc.get("total_annual_cost", 0)
            doc.add_paragraph(f"Estimated Annual Cost: £{cost:,}")

        roi = fin.get("roi_analysis") or {}
        if roi:
            doc.add_heading("Return on Investment", level=2)
            doc.add_paragraph(
                f"Remediation Investment: £{roi.get('remediation_investment', 0):,}"
            )
            doc.add_paragraph(
                f"Payback Period: {roi.get('payback_months', '—')} months"
            )
            doc.add_paragraph(
                f"3-Year Net Benefit: £{roi.get('roi_year3', 0):,}"
            )

    def _add_appendix(self, doc: Document) -> None:
        doc.add_heading("Appendix", level=1)

        doc.add_heading("Assessment Methodology", level=2)
        doc.add_paragraph(
            "This assessment was conducted using 0xRecon's automated "
            "cryptographic asset discovery and policy evaluation engine. "
            "Findings are generated by comparing discovered assets against "
            "the configured security policy rules."
        )

        doc.add_heading("Scan Information", level=2)
        doc.add_paragraph(f"Scan Name: {self._data.get('scan_name', '—')}")
        doc.add_paragraph(f"Generated: {self._data.get('generated_at', '—')}")

        doc.add_heading("Disclaimer", level=2)
        doc.add_paragraph(
            "Financial estimates are based on industry benchmarks and "
            "statistical models. Actual costs may vary based on organisation "
            "size, sector, and specific risk factors. All values in GBP."
        )
